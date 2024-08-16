import re
import time
from pathlib import Path

from PIL import Image
from docx import Document
from docx.image.exceptions import UnrecognizedImageError

from src.api.dify_api import DifyApi
from src.database.ai_database import AiDatabase
from src.database.dify_database import DifyDatabase
from src.utils.config import config


def get_knowledge_base_documents():
    source_url = config.mapping['url']['source']
    source_api = DifyApi(source_url, config.mapping['secret_key']['source'])
    ai_database = AiDatabase()
    for item in config.mapping['knowledge_base']:
        source_dataset_id = source_api.get_dataset_id_by_name(item['source'])
        source_knowledge_base = {'id': source_dataset_id, 'url': source_url, 'name': item['source']}
        ai_database.save_knowledge_base_info(source_knowledge_base)
        source_documents = source_api.get_documents_in_dataset(source_dataset_id)
        ai_database.delete_no_exist_documents(source_dataset_id, source_documents)
        for document in source_documents:
            document['dataset_id'] = source_dataset_id
            ai_database.save_document(document)
            segments = source_api.get_segments_from_document(source_dataset_id, document['id'])
            ai_database.delete_no_exist_segment(document['id'], [segment['id'] for segment in segments])
            for segment in segments:
                keywords = segment['keywords']
                keywords.sort()
                segment['keywords'] = ','.join(keywords)
                ai_database.save_document_segment(segment)


def get_document_id_by_name(name, documents):
    for document in documents:
        if document['name'].rpartition('.')[0].lower() == name.lower():
            return document['id']
    return None


def upload_document_to_knowledge_base(api, knowledge_base, documents: list, is_replace_document=True):
    dataset_id = api.get_dataset_id_by_name(knowledge_base)
    exist_documents = api.get_documents_in_dataset(dataset_id)
    for document in documents:
        document_name_without_extension = document['name'].rpartition('.')[0]
        if is_replace_document:
            document_id = get_document_id_by_name(document_name_without_extension, exist_documents)
            if document_id is not None:
                api.delete_document(dataset_id, document_id)
        document_id = api.create_document(dataset_id, document_name_without_extension)
        sorted_segments = sorted(document['segment'], key=lambda x: x['position'])
        for segment in sorted_segments:
            api.create_segment_in_document(dataset_id, document_id,
                                           segment['content'], segment['answer'], segment['keywords'])


def add_documents_to_knowledge_base():
    source_url = config.mapping['url']['source']
    target_url = config.mapping['url']['target']
    target_api = DifyApi(target_url, config.mapping['secret_key']['target'])
    ai_database = AiDatabase()
    for item in config.mapping['knowledge_base']:
        documents = ai_database.get_knowledge_base_documents(source_url, item['source'])
        upload_document_to_knowledge_base(target_api, item['target'], documents)


def get_images_from_document(document_segments: list) -> list:
    pattern = r'!\[image\]\(/files/(.*?)/image-preview\)'
    processed_data = []
    for record in document_segments:
        images = []
        for segment in record['segments']:
            uuids = re.findall(pattern, segment['content'])
            images.extend(uuids)
        if images:
            new_record = {
                'dataset_id': record['dataset_id'],
                'document_id': record['document_id'],
                'images': images
            }
            processed_data.append(new_record)
    return processed_data


def convert_image_to_jpg(image_path: Path) -> Path:
    jpg_image_path = f'{image_path.parent.parent / Path("converted") / image_path.stem}.jpg'
    Image.open(image_path).convert('RGB').save(jpg_image_path)
    return Path(jpg_image_path)


def add_images_to_word_file(images: list, file_path: Path):
    doc = Document()
    for image in images:
        try:
            doc.add_picture(image.as_posix())
        except UnrecognizedImageError:
            jpg_image_path = convert_image_to_jpg(image)
            doc.add_picture(jpg_image_path.as_posix())
        doc.add_paragraph()
    doc.save(file_path.as_posix())


def get_images_from_segments(data_list: list):
    pattern = r'!\[image\]\(/files/(.*?)/image-preview\)'
    images = []
    for record in data_list:
        images.extend(re.findall(pattern, record['content']))
    return images


class IndexingNotCompletedError(Exception):
    pass


def replace_images_in_documents():
    source_url = config.mapping['url']['source']
    source_api = DifyApi(source_url, config.mapping['secret_key']['source'])
    target_url = config.mapping['url']['target']
    target_api = DifyApi(target_url, config.mapping['secret_key']['target'])
    assets_file_root_path = Path(__file__).parent.absolute() / Path('src/assets')
    word_file_root_path = assets_file_root_path / Path('word_files')

    for dataset in config.mapping['knowledge_base']:
        documents_list = []
        source_dataset_id = source_api.get_dataset_id_by_name(dataset['source'])
        source_documents = source_api.get_documents_in_dataset(source_dataset_id)
        for document in source_documents:
            segments = source_api.get_segments_from_document(source_dataset_id, document['id'])
            document_segments = {
                'dataset_id': source_dataset_id,
                'document_id': document['id'],
                'segments': [
                    {
                        'id': segment['id'],
                        'position': segment['position'],
                        'content': segment['content'],
                    }
                    for segment in segments
                ]
            }
            documents_list.append(document_segments)
        if documents_list:
            documents_with_images = get_images_from_document(documents_list)
            if documents_with_images:
                images_mapping = {}
                target_dataset_id = target_api.get_dataset_id_by_name(dataset['target'])
                dify_database = DifyDatabase()
                for document in documents_with_images:
                    source_images = document['images']
                    images = []
                    for id_ in source_images:
                        image_path = assets_file_root_path / dify_database.get_image_path(id_)
                        images.append(image_path)
                    word_file_path = word_file_root_path / f'{document["document_id"]}.docx'
                    add_images_to_word_file(images, word_file_path)
                    response = target_api.create_document_by_file(target_dataset_id, word_file_path)
                    images_document_id = response['document']['id']
                    target_images_segments = target_api.get_segments_from_document(target_dataset_id,
                                                                                   images_document_id)
                    limit = 0
                    while target_api.get_document_embedding_status(target_dataset_id, response['batch'],
                                                                   images_document_id) != 'completed' and limit < 3:
                        limit += 1
                        time.sleep(5)
                    if target_api.get_document_embedding_status(target_dataset_id, response['batch'],
                                                                images_document_id) != 'completed' and limit == 3:
                        raise IndexingNotCompletedError('Indexing not completed')
                    target_api.delete_document(target_dataset_id, images_document_id)
                    target_images = get_images_from_segments(target_images_segments)
                    images_mapping.update(dict(zip(source_images, target_images)))

                target_documents = target_api.get_documents_in_dataset(target_dataset_id)
                for document in target_documents:
                    segments = source_api.get_segments_from_document(target_dataset_id, document['id'])
                    for segment in segments:
                        origin_segment_content = segment['content']
                        for key, value in images_mapping.items():
                            segment['content'] = segment['content'].replace(key, value)
                        if segment['content'] != origin_segment_content:
                            target_api.update_segment_in_document(
                                target_dataset_id, segment['document_id'], segment['id'],
                                segment['content'], segment['answer'], segment['keywords'], True)


def main():
    get_knowledge_base_documents()
    add_documents_to_knowledge_base()
    replace_images_in_documents()


if __name__ == '__main__':
    main()
