import random
import re
import time
import uuid
from pathlib import Path
from typing import Union, Optional

import pandas as pd
from PIL import Image
from docx import Document
from docx.image.exceptions import UnrecognizedImageError

from src.api.dataset_api import DatasetApi
from src.database.dify_database import DifyDatabase
from src.database.record_database import RecordDatabase
from src.utils.config import config
from src.utils.document_sync_config import DocumentSyncConfig
from src.utils.time_utils import timing


class IndexingNotCompletedError(Exception):
    pass


class SplitCountExceeded(Exception):
    pass


class KnowledgeBase(object):
    def __init__(self, env, dataset_id, dataset_name, api: DatasetApi,
                 db: DifyDatabase = None, record_db: RecordDatabase = None):
        self.env = env
        self.dataset_id = dataset_id
        self.dataset_name = dataset_name
        self.api = api
        self.db = db
        self.record_db = record_db

    def record_knowledge_base_info(self):
        if not self.record_db:
            raise Exception('Record database is not set')
        knowledge_base_info = {'id': self.dataset_id, 'url': self.api.base_url, 'name': self.dataset_name}
        self.record_db.save_knowledge_base_info(knowledge_base_info)

    def record_documents(self, documents):
        if not self.record_db:
            raise Exception('Record database is not set')
        docs_in_record = self.fetch_documents(source='record', with_segment=True)
        if docs_in_record is not None:
            docs_to_remove_in_record = [doc for doc in docs_in_record if
                                        doc['id'] not in [document['id'] for document in documents]]
            if docs_to_remove_in_record:
                self.record_db.remove_documents([doc['id'] for doc in docs_to_remove_in_record])
        modified_documents = [
            {k: (self.dataset_id if k == 'dataset_id' else v) for k, v in document.items() if k != 'segment'}
            for document in documents
        ]
        self.record_db.save_documents(modified_documents)
        for document in documents:
            if 'segment' not in document:
                continue
            segment_ids = [segment['id'] for segment in document['segment']]
            segment_ids_in_record = [segment['id'] for segment in self._fetch_segments('record', document['id'])]
            segments_to_remove_in_record = [segment_id for segment_id in segment_ids_in_record if
                                            segment_id not in segment_ids]
            if segments_to_remove_in_record:
                self.record_db.remove_segments(document['id'], segments_to_remove_in_record)
            for segment in document['segment']:
                keywords = segment['keywords']
                if isinstance(keywords, list):
                    keywords.sort()
                    segment['keywords'] = ','.join(keywords)
            self.record_db.save_segments(document['segment'])

    def _fetch_all_documents(self, source, is_enabled: bool):
        if source == 'api':
            return self.api.get_documents_in_dataset(self.dataset_id, is_enabled)
        elif source == 'db':
            if self.db is None:
                raise Exception('Dify database is not set')
            return self.db.get_documents(self.dataset_id, with_segment=False, is_enabled=is_enabled)
        elif source == 'record':
            return self.record_db.get_documents(
                self.api.base_url, self.dataset_id, with_segment=False, is_enabled=is_enabled
            )
        else:
            return None

    def _find_document_by_id(self, documents, document_id):
        for document in documents:
            if document['id'] == document_id:
                return document
        return None

    def _get_images_from_segments(self, segments):
        pattern = r'(?:!\[image\])?\([^)]*/files/(.*?)/(?:image-preview|file-preview)\)'
        images = []
        for segment in segments:
            uuids = re.findall(pattern, segment['content'])
            images.extend(uuids)
        return images

    def _fetch_segments(self, source, document_id):
        if source == 'api':
            return self.api.get_segments_from_document(self.dataset_id, document_id)
        elif source == 'db':
            if self.db is None:
                raise Exception('Dify database is not set')
            return self.db.get_segments(document_id)
        elif source == 'record':
            return self.record_db.get_segments(document_id)
        return None

    def _process_document(self, document, source, with_segment, with_image):
        document['dataset_id'] = self.dataset_id
        if with_segment:
            document['segment'] = self._fetch_segments(source, document['id'])
        if with_image and 'segment' in document:
            document['image'] = self._get_images_from_segments(document['segment'])

    def fetch_documents(self, source, document_id=None, with_segment=False, with_image=False,
                        is_enabled: bool = None) -> Optional[Union[dict, list[dict]]]:
        documents = self._fetch_all_documents(source, is_enabled)
        if documents is None or not documents:
            return None

        if document_id:
            document = self._find_document_by_id(documents, document_id)
            if document:
                self._process_document(document, source, with_segment, with_image)
                return document
            return None
        else:
            for document in documents:
                self._process_document(document, source, with_segment, with_image)
            return documents

    def get_document_id_by_name(self, name, documents):
        if documents is not None:
            for document in documents:
                if document['name'] == name:
                    return document['id']
        return None

    @timing
    def sync_documents(self, documents: Union[dict, list[dict]], sync_config: DocumentSyncConfig,
                       source: str = 'api') -> dict:
        documents = self._handle_and_sort_text(documents, sync_config.preserve_document_order)
        for doc in documents:
            if 'segment' in doc:
                doc['segment'] = self._handle_and_sort_text(doc['segment'], sync_config.preserve_segment_order)

        current_documents, existing_ids, extra_ids = self._fetch_and_filter_current_documents(documents)

        if sync_config.skip_existing:
            existing_doc_names = {doc['name'] for doc in current_documents if doc['id'] in existing_ids}
            documents = [doc for doc in documents if doc['name'] not in existing_doc_names]
            for doc_name in existing_doc_names:
                print(f'Skip existing document: {doc_name}')
        else:
            if sync_config.replace_existing and existing_ids:
                if sync_config.backup:
                    self.backup_documents(document_ids=existing_ids, source=source)
                self.delete_documents(existing_ids)
        if sync_config.remove_extra and extra_ids:
            if sync_config.backup:
                self.backup_documents(document_ids=existing_ids, source=source)
            self.delete_documents(extra_ids)

        return self.create_document_by_text(documents)

    def _wait_document_embedding(self, batch_id, document_id, status='completed', retry: int = 600):
        index = 0
        while self.api.get_document_embedding_status(self.dataset_id, batch_id, document_id) != status:
            if index == retry:
                raise IndexingNotCompletedError(
                    f'Indexing not completed after {retry} attempts for document_id: {document_id}')
            index += 1
            time.sleep(0.5)

    def create_document_by_file(self, file_path):
        response = self.api.create_document_by_file(self.dataset_id, file_path)
        if response.data is None:
            return None
        document_id = response.data['document']['id']
        batch_id = response.data['batch']
        self._wait_document_embedding(batch_id, document_id)
        return document_id

    def delete_documents(self, document_ids: list[str]):
        for document_id in document_ids:
            if document_id:
                self.api.delete_document(self.dataset_id, document_id)

    def update_segment_in_document(self, segment):
        self.api.update_segment_in_document(
            self.dataset_id,
            segment['document_id'],
            segment['id'],
            segment['content'],
            segment['answer'],
            segment['keywords'],
            segment.get('enabled', True)
        )

    def empty_dataset(self):
        documents = self.api.get_documents_in_dataset(self.dataset_id)
        for document in documents:
            self.api.delete_document(self.dataset_id, document['id'])

    def upload_images(self, images_path: list, doc_name: str = uuid.uuid4()) -> dict:
        images_mapping = {}

        docs_with_images = self.upload_images_by_word_file(doc_name, images_path)
        images = []
        for document_id in docs_with_images:
            if document_id == '':
                images.append('')
            else:
                document = self.fetch_documents('api', document_id=document_id, with_segment=True, with_image=True)
                images.extend(document['image'])
        images_mapping.update(dict(zip(images_path, images)))
        self.delete_documents(docs_with_images)

        return images_mapping

    def upload_images_by_word_file(self, document_name, images, split_count=1, current_split=1, max_split=5) -> list:
        def split_list(lst, count):
            k, m = divmod(len(lst), count)
            return [lst[i * k + min(i, m):(i + 1) * k + min(i + 1, m)] for i in range(count)]

        if current_split > max_split:
            raise SplitCountExceeded(
                f'Max split count of {max_split} exceeded for document {document_name} uploaded to {self.dataset_name}'
            )

        uploaded_documents = []
        split_image_paths = split_list(images, split_count)
        for batch_index, image_batch in enumerate(split_image_paths):
            if image_batch:
                word_file_path = config.word_dir_path / Path(f'{document_name}-{current_split}-{batch_index}.docx')
                self.add_images_to_word_file(image_batch, word_file_path)
                document_id = self.create_document_by_file(word_file_path)
                if document_id is not None:
                    uploaded_documents.append(document_id)
                else:
                    if len(image_batch) == 1:
                        uploaded_documents.append('')
                    else:
                        uploaded_documents.extend(
                            self.upload_images_by_word_file(
                                document_name, image_batch, split_count * 2, current_split + 1, max_split
                            )
                        )
            time.sleep(random.uniform(1, 3))
        return uploaded_documents

    def add_images_to_word_file(self, images: list[Path], word_file: Path):
        doc = Document()
        for image in images:
            try:
                doc.add_picture(image.as_posix())
            except UnrecognizedImageError:
                jpg_image_path = self.convert_image_to_jpg(image)
                doc.add_picture(jpg_image_path.as_posix())
            doc.add_paragraph()
        doc.save(word_file.as_posix())

    def convert_image_to_jpg(self, image_path: Path) -> Path:
        jpg_image_path = config.convert_dir_path / Path(f'{image_path.stem}.jpg')
        Image.open(image_path).convert('RGB').save(jpg_image_path)
        return Path(jpg_image_path)

    def _handle_and_sort_text(self, text: Union[dict, list[dict]], sort_text: bool) -> list[dict]:
        if isinstance(text, dict):
            text = [text]
        elif not isinstance(text, list):
            raise ValueError("The text must be either a dict or a list of dicts")
        if sort_text and all('position' in item for item in text):
            return sorted(text, key=lambda x: x['position'])
        return text

    def _fetch_and_filter_current_documents(self, documents):
        current_documents = self.fetch_documents(source='db') or []
        document_names = {document['name'] for document in documents}
        existing_ids = [doc['id'] for doc in current_documents if doc['name'] in document_names]
        extra_ids = [doc['id'] for doc in current_documents if doc['name'] not in document_names]
        return current_documents, existing_ids, extra_ids

    def create_document_by_text(self, documents: list[dict]) -> dict:
        docs_name_id_mapping = {}
        for document in documents:
            document_id, batch_id = self.api.create_document(self.dataset_id, document['name'])
            self._wait_document_embedding(batch_id, document_id)
            if 'segment' in document:
                segments = document['segment']
                for segment in segments:
                    segment_id = self.api.create_segment_in_document(self.dataset_id, document_id, segment)
                    if not segment.get('enabled'):
                        self.api.update_segment_in_document(
                            self.dataset_id, document_id, segment_id, segment.get('content'), segment.get('answer'),
                            segment.get('keywords'), False
                        )
            docs_name_id_mapping[document['name']] = document_id
        return docs_name_id_mapping

    def get_image_paths(self, image_uuids: list[str]):
        image_paths = {}

        for image_uuid in image_uuids:
            image_path = self.db.get_image_path(image_uuid)
            if image_path is None:
                continue
            image_paths[image_uuid] = image_path
        return image_paths

    def backup_documents(self, document_ids: list[str], source: str = 'api'):
        documents = []
        for doc_id in document_ids:
            document = self.fetch_documents(source=source, document_id=doc_id, with_segment=True)
            if document is None:
                continue
            documents.append(document)
        rows = []
        for doc in documents:
            for segment in doc['segment']:
                rows.append({
                    'document_name': doc['name'],
                    'segment_position': segment['position'],
                    'content': segment['content'],
                    'answer': segment['answer'],
                    'keywords': segment['keywords'],
                })
        documents_df = pd.DataFrame(rows)
        documents_df['environment'] = self.env
        documents_df['dataset_name'] = self.dataset_name
        self.record_db.backup_documents(documents=documents_df)

    def disable_documents(self, document_ids: list[str], source: str = 'api'):
        documents = []
        for doc_id in document_ids:
            document = self.fetch_documents(source=source, document_id=doc_id, with_segment=True)
            if document is None:
                continue
            documents.append(document)
        for doc in documents:
            for segment in doc['segment']:
                if not segment.get('enabled'):
                    continue
                segment['enabled'] = False
                self.update_segment_in_document(segment)
            print(f'Disabled document: {doc["name"]}')
