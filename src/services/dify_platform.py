import json
import random
import re
import time
import uuid
from pathlib import Path

from PIL import Image
from docx import Document
from docx.image.exceptions import UnrecognizedImageError

from src.api.app_api import AppApi
from src.api.dataset_api import DatasetApi
from src.database.dify_database import DifyDatabase
from src.database.record_database import RecordDatabase
from src.services.knowledge_base import KnowledgeBase
from src.services.s3_handler import S3Handler
from src.utils.config import config


class SplitCountExceeded(Exception):
    pass


class DifyPlatform(object):
    def __init__(self, api_config, record_db_name: str = 'record', s3_config=None, dify_db_name: str = None):
        self.app_api = AppApi(api_config.url, api_config.app_token)
        self.dataset_api = DatasetApi(api_config.url, api_config.dataset_token)
        self.datasets = self.dataset_api.get_datasets()
        self.record_db = RecordDatabase(record_db_name)
        self._s3_config = s3_config
        self._s3 = None
        self._dify_db_name = dify_db_name
        self._dify_db = None

    @property
    def s3(self):
        if self._s3 is None:
            if self._s3_config is None:
                raise Exception("'s3_config' is not set, provide valid 's3_config'")
            self._s3 = S3Handler(
                self._s3_config.access_key_id,
                self._s3_config.secret_access_key,
                self._s3_config.region,
                self._s3_config.bucket
            )
        return self._s3

    @property
    def dify_db(self):
        if self._dify_db is None:
            if self._dify_db_name is None:
                raise Exception("'dify_db_name' is not set, provide valid 'dify_db_name'")
            self._dify_db = DifyDatabase(self._dify_db_name)
        return self._dify_db

    def get_dataset_id_by_name(self, name) -> str:
        dataset_id = None
        for dataset in self.datasets:
            if dataset['name'] == name:
                dataset_id = dataset['id']
                break
        return dataset_id

    def download_images_to_local(self, image_uuids: list):
        image_paths = {}
        for uuid in image_uuids:
            image_path_in_dify = self.dify_db.get_image_path(uuid)
            if self.s3.find_and_download_file(str(image_path_in_dify.as_posix()), config.image_dir_path):
                image_path = config.image_dir_path / image_path_in_dify.name
                image_paths[uuid] = image_path
        return image_paths

    def convert_image_to_jpg(self, image_path: Path) -> Path:
        jpg_image_path = config.convert_dir_path / Path(f'{image_path.stem}.jpg')
        Image.open(image_path).convert('RGB').save(jpg_image_path)
        return Path(jpg_image_path)

    def add_images_to_document(self, images: list[Path], word_file: Path):
        doc = Document()
        for image in images:
            try:
                doc.add_picture(image.as_posix())
            except UnrecognizedImageError:
                jpg_image_path = self.convert_image_to_jpg(image)
                doc.add_picture(jpg_image_path.as_posix())
            doc.add_paragraph()
        doc.save(word_file.as_posix())

    def upload_images_as_word_documents(self, dataset, document_name, images, split_count=1,
                                        current_split=1, max_split=5) -> list:
        def split_list(lst, count):
            k, m = divmod(len(lst), count)
            return [lst[i * k + min(i, m):(i + 1) * k + min(i + 1, m)] for i in range(count)]

        if current_split > max_split:
            raise SplitCountExceeded(
                f'Max split count of {max_split} exceeded for document {document_name} uploaded to {dataset.dataset_id}'
            )

        uploaded_documents = []
        split_image_paths = split_list(images, split_count)
        for batch_index, image_batch in enumerate(split_image_paths):
            if image_batch:
                word_file_path = config.word_dir_path / Path(f'{document_name}-{current_split}-{batch_index}.docx')
                self.add_images_to_document(image_batch, word_file_path)
                document_id = dataset.create_document_by_file(word_file_path)
                if document_id is not None:
                    uploaded_documents.append(document_id)
                else:
                    if len(image_batch) == 1:
                        uploaded_documents.append('')
                    else:
                        uploaded_documents.extend(
                            self.upload_images_as_word_documents(
                                dataset, document_name, image_batch, split_count * 2, current_split + 1, max_split
                            )
                        )
            time.sleep(random.uniform(1, 3))
        return uploaded_documents

    def upload_images_to_dify(self, images_path: list, dataset: KnowledgeBase, doc_name: str = uuid.uuid4()) -> dict:
        images_mapping = {}

        docs_with_images = self.upload_images_as_word_documents(dataset, doc_name, images_path)
        images = []
        for document_id in docs_with_images:
            if document_id == '':
                images.append('')
            else:
                document = dataset.get_documents('api', document_id=document_id, with_segment=True, with_image=True)
                images.extend(document['image'])
        images_mapping.update(dict(zip(images_path, images)))
        dataset.delete_document(docs_with_images)

        return images_mapping

    def fix_json_str(self, json_str):
        json_str = re.sub(r'^[^{]*', '', json_str)
        json_str = re.sub(r'\s*[^}\n]*$', '', json_str)
        last_quote_index = json_str.rfind('"')
        last_right_square_index = json_str.rfind(']')
        last_brace_index = json_str.rfind('}')
        if (last_quote_index > last_right_square_index
                and re.search(r'[^\s\n]', json_str[last_quote_index + 1:last_brace_index])):
            json_str = json_str[:last_brace_index] + '"' + json_str[last_brace_index:]
        return json_str

    def analyze_content(self, query: str):
        response = self.app_api.query_ai_agent(query, response_mode='streaming')
        try:
            answer = json.loads(self.fix_json_str(response.get('answer', {})))
        except json.JSONDecodeError:
            answer = {}
        return answer

    def init_knowledge_base(self, dataset_name):
        dataset_id = self.get_dataset_id_by_name(dataset_name)
        return KnowledgeBase(self.dataset_api, dataset_id, dataset_name, self.record_db)
