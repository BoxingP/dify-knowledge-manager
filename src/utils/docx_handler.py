import base64
import re
from collections import namedtuple
from typing import Any, Optional
from xml.etree import ElementTree

import pandas as pd
from docx import Document
from docx.oxml import CT_P, CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.services.knowledge_base import KnowledgeBase
from src.utils.config import config


class DocxHandler(object):
    def __init__(self, file_path, title_prefix: str = '#', text_rules: Optional[dict[str, str]] = None):
        self.file_path = file_path
        self.document = Document(self.file_path)
        self.title_prefix = title_prefix
        self.text_rules = text_rules or {}

    def _extract_block_items(self):
        for child in self.document.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, self.document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, self.document)

    def _extract_image_data(self, paragraph: Paragraph, image_counter: int) -> Optional[dict[str, Any]]:
        img_data = None
        for run in paragraph.runs:
            xml_str = run.element.xml
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            }
            root = ElementTree.fromstring(xml_str)

            for pic in root.findall('.//pic:pic', namespaces):
                cnvpr_elem = pic.find('pic:nvPicPr/pic:cNvPr', namespaces)
                blip_elem = pic.find('pic:blipFill/a:blip', namespaces)

                if cnvpr_elem is None or blip_elem is None:
                    continue

                name = cnvpr_elem.get('name', f'image_{image_counter}')
                embed = blip_elem.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )

                if embed:
                    image_part = self.document.part.related_parts.get(embed)
                    if image_part:
                        return {
                            'image_id': image_counter,
                            'image_name': name,
                            'image_type': image_part.content_type.split('/')[-1],
                            'image_base64_string': base64.b64encode(image_part._blob).decode()
                        }

        return img_data

    def _process_images(self, images: pd.DataFrame, image_reference_type: str = 'dify',
                        knowledge_base: KnowledgeBase = None) -> dict:
        if image_reference_type == 'dify' and knowledge_base is None:
            return {}

        image_paths = []
        for _, image in images.iterrows():
            img_path = config.image_dir_path / f"{self.file_path.stem}.{image['image_id']}.{image['image_name']}.{image['image_type']}"
            with open(img_path, 'wb') as f:
                f.write(base64.b64decode(image['image_base64_string']))
            image_paths.append(img_path)

        uploaded = knowledge_base.upload_images(image_paths, self.file_path.stem)
        return {
            str(index): f'\n![image](/files/{value}/file-preview)\n'
            for index, (key, value) in enumerate(uploaded.items())
            if value
        }

    def extract_content(self) -> namedtuple:
        Content = namedtuple('Content', ['document', 'images', 'tables'])

        doc_rows = []
        img_rows = []
        tbl_rows = []

        image_counter = 0

        for block in self._extract_block_items():
            if isinstance(block, Paragraph):
                img_data = self._extract_image_data(block, image_counter)
                if img_data:
                    img_rows.append(img_data)
                    doc_rows.append({
                        'text': '',
                        'style': '',
                        'image_id': image_counter,
                        'table_id': None,
                    })
                    image_counter += 1
                else:
                    doc_rows.append({
                        'text': block.text.strip(),
                        'style': block.style.name if block.style else '',
                        'image_id': None,
                        'table_id': None,
                    })
            elif isinstance(block, Table):
                table_id = len(tbl_rows)
                tbl_rows.append({
                    'table_id': table_id,
                    'content': pd.DataFrame([[cell.text for cell in row.cells] for row in block.rows]).to_markdown()
                })
                doc_rows.append({
                    'text': '',
                    'style': '',
                    'image_id': None,
                    'table_id': table_id,
                })

        document_df = pd.DataFrame(doc_rows)
        document_df[['image_id', 'table_id']] = document_df[['image_id', 'table_id']].applymap(
            lambda x: '' if pd.isnull(x) else str(int(x))
        )
        return Content(
            document=document_df,
            images=pd.DataFrame(img_rows),
            tables=pd.DataFrame(tbl_rows)
        )

    def _apply_text_rules(self, text: str) -> str:
        for pattern, replacement in self.text_rules.items():
            text = re.sub(pattern, replacement, text)
        return text

    def convert_to_str(self, content: namedtuple, **kwargs) -> str:
        images_dict = self._process_images(content.images, **kwargs) if not content.images.empty else {}
        tables_dict = {
            str(table['table_id']): table['content'] for _, table in content.tables.iterrows()
        } if not content.tables.empty else {}

        def process_row(row):
            text = row.get('text', '').strip()
            if text:
                text = self._apply_text_rules(text)
                if row.get('style', '').lower() == 'title':
                    return f'{self.title_prefix} {text}'
                return text

            if row.get('image_id', '').strip():
                return images_dict.get(row['image_id'], '[image]')

            if row.get('table_id', '').strip():
                return tables_dict.get(row['table_id'], '[table]')

            return ''

        return '\n'.join(filter(None, (process_row(row) for _, row in content.document.iterrows())))
